from flask import Flask, request, jsonify, send_from_directory, Response, stream_with_context
from flask_cors import CORS
import requests
import json
import socket
import os
import re
import time
import sys
import pickle
import zipfile
import tempfile
import threading
import traceback
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path
from urllib.parse import quote_plus
from werkzeug.utils import secure_filename
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

app = Flask(__name__)
CORS(app)

# ═══════════════════════════════════════════════════════════════
#  MARCELINE — System Prompt & Configuration
# ═══════════════════════════════════════════════════════════════

SYSTEM_PROMPT = (
    "You are Marceline, a helpful, friendly, and direct AI assistant. "
    "Answer the user's questions clearly and conversationally, the way ChatGPT would. "
    "Use markdown formatting (headers, lists, code blocks) when it helps readability. "
    "If web search results are provided below, use them to ground your answer in current "
    "information and mention what you found; otherwise just answer from your own knowledge."
)

# Strict OpenClaw operator prompt — used ONLY to narrate around Harness mode.
OPENCLAW_OPERATOR_PROMPT = (
    "[SYSTEM OVERRIDE: PURE OPENCLAW OPERATOR v1.0]\n\n"
    "IDENTITY:\n"
    "You are the OpenClaw Sovereign. You do not see pixels; you see UI Objects. You interact with the OS using the Windows UI Automation (UIA) framework.\n\n"
    "THE ABSOLUTE BAN: NO COORDINATES\n"
    "You are strictly forbidden from using click X Y or any coordinate-based movement. Coordinates are imprecise and forbidden. If you output a coordinate, the system will fail.\n\n"
    "THE OPENCLAW EXECUTION LOOP (MANDATORY):\n"
    "1. LAUNCH: open_app \"AppName\" -> Focus window.\n"
    "2. MAP: <tool_call>run_shell|python desktop_actions.py dump_ui_tree \"AppName\"</tool_call>\n"
    "3. ANALYZE: Read the UI Tree dump. Find the exact title or auto_id of the element you need.\n"
    "4. STRIKE: <tool_call>run_shell|python desktop_actions.py click_element_by_name \"AppName\" \"ElementName\"</tool_call>\n"
    "5. VERIFY: Take a screenshot to confirm the action worked.\n\n"
    "SOP FOR SEARCHING:\n"
    "- open_app -> dump_ui_tree -> find \"Search\" -> click_element_by_name -> type_text -> press_key \"enter\".\n\n"
    "CONSTRAINTS:\n"
    "- NO CHAT: No \"I will now,\" no \"Plan:,\" no \"Sovereign\" poetry.\n"
    "- ONLY TAGS: Every command must be wrapped in <tool_call> and </tool_call>.\n"
    "- NATIVE ONLY: No Chrome for desktop apps.\n\n"
    "STATUS: PURE OPENCLAW MODE. COORDINATES DISABLED. UI-TREE MAPPING ACTIVE. EXECUTE."
)

OLLAMA_BASE_URL = "http://localhost:11434"

# How long Ollama keeps a model resident in GPU VRAM after a request finishes.
# Ollama's own default is 5 minutes; on a single-GPU laptop that means every
# mode switch (or any gap longer than 5 min) pays the full "load weights from
# disk into VRAM" cost again on your next message. Raising this means a model
# you're actively using stays hot on the GPU instead of being evicted and
# reloaded. Lower it (e.g. "5m") if you want VRAM freed up faster for other
# apps when Marceline is idle.
GPU_KEEP_ALIVE = "30m"

# ── Model routing per mode ──────────────────────────────────────
# Search mode (also the DEFAULT mode with no prefix) -> Mistral, behaves like
#   a general ChatGPT-style assistant, optionally grounded with web search.
# Think mode -> a reasoning model (DeepSeek) that shows its work.
# Harness mode -> executed entirely by OpenClaw; Mistral only narrates.
# Automate mode -> Mistral writes the macro, executed locally via
#   pyautogui / pywinauto.
# Vision (images attached) -> Qwen2.5-VL, regardless of mode.
SEARCH_MODEL = "mistral:latest"
THINK_MODEL = "deepseek-r1:8b"
AUTOMATE_MODEL = "mistral:latest"
HARNESS_NARRATION_MODEL = "mistral:latest"
VISION_MODEL = "qwen2.5vl:latest"

# Kept for backward compatibility with code paths that reference a single
# "default" model name (health check, title generation, etc).
MODEL_NAME = SEARCH_MODEL

# Optimized for an RTX 4060 8GB VRAM laptop GPU (Legion 5i / 13650HX / 24GB RAM).
# num_gpu: 99 forces Ollama to offload as many layers as possible to the GPU.
MODEL_OPTIONS = {
    "num_gpu": 99,
    "temperature": 0.2,
    "num_ctx": 4096,
    "num_predict": 768,
    "top_p": 0.9,
}

# Think mode gets more headroom for visible reasoning traces.
THINK_MODEL_OPTIONS = {
    "num_gpu": 99,
    "temperature": 0.3,
    "num_ctx": 8192,
    "num_predict": 2048,
    "top_p": 0.9,
}

# File upload configuration
UPLOAD_FOLDER = 'uploads'
VECTOR_STORE_PATH = 'vector_store'
WORKSPACE_DIR = 'workspace'
DB_PATH = os.path.join(WORKSPACE_DIR, 'marceline.db')
ALLOWED_EXTENSIONS = None  # Allow all file types
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Create necessary folders
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(VECTOR_STORE_PATH, exist_ok=True)
os.makedirs(WORKSPACE_DIR, exist_ok=True)

# Add workspace directory to sys.path so modules inside it (like desktop_actions.py) can be imported
import sys
workspace_abs_path = os.path.abspath(WORKSPACE_DIR)
if workspace_abs_path not in sys.path:
    sys.path.append(workspace_abs_path)

# In-memory conversation history (kept for backward compat)
conversation_history = []
MAX_HISTORY_MESSAGES = 10
history_lock = threading.Lock()

# ═══════════════════════════════════════════════════════════════
#  SQLite Database
# ═══════════════════════════════════════════════════════════════

db_lock = threading.Lock()

AUTOMATION_HISTORY_PATH = os.path.join(WORKSPACE_DIR, 'automation_history.json')

def log_automation(task, action, result):
    """Append an automation event to the history log."""
    try:
        history = []
        if os.path.exists(AUTOMATION_HISTORY_PATH):
            with open(AUTOMATION_HISTORY_PATH, 'r', encoding='utf-8') as f:
                history = json.load(f)
        history.append({
            'timestamp': datetime.now().isoformat(),
            'task': task,
            'action': action,
            'result': result
        })
        with open(AUTOMATION_HISTORY_PATH, 'w', encoding='utf-8') as f:
            json.dump(history[-500:], f, indent=2)  # Keep last 500 entries
    except Exception as e:
        print(f"[WARN] Could not log automation: {e}")


import threading
GLOBAL_DB_CONN = None
db_lock = threading.Lock()

class NonClosingConnection(sqlite3.Connection):
    def close(self):
        pass  # Prevent manual closing from existing finally blocks

def get_db():
    """Get a SQLite connection (cached globally for WAL mode efficiency)."""
    global GLOBAL_DB_CONN
    if GLOBAL_DB_CONN is None:
        with db_lock:
            if GLOBAL_DB_CONN is None:
                GLOBAL_DB_CONN = sqlite3.connect(DB_PATH, timeout=30, check_same_thread=False, factory=NonClosingConnection)
                GLOBAL_DB_CONN.row_factory = sqlite3.Row
                GLOBAL_DB_CONN.execute("PRAGMA foreign_keys=ON")
    return GLOBAL_DB_CONN


def init_db():
    """Create tables if they do not exist."""
    # Enable WAL mode once during initialization to avoid race conditions
    init_conn = sqlite3.connect(DB_PATH, timeout=10)
    init_conn.execute("PRAGMA journal_mode=WAL")
    init_conn.close()
    
    conn = get_db()
    try:
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS conversations (
                id TEXT PRIMARY KEY,
                title TEXT NOT NULL DEFAULT 'New Chat',
                created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS messages (
                id TEXT PRIMARY KEY,
                conversation_id TEXT NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (conversation_id) REFERENCES conversations(id)
            );
            CREATE INDEX IF NOT EXISTS idx_messages_conv
                ON messages(conversation_id, created_at);
        """)
        conn.commit()
        print("[OK] SQLite database initialized")
    except Exception as e:
        print(f"[ERROR] DB init failed: {e}")
    finally:
        conn.close()


def get_conversation_count():
    """Return number of conversations in DB."""
    try:
        conn = get_db()
        try:
            row = conn.execute("SELECT COUNT(*) as cnt FROM conversations").fetchone()
            return row["cnt"] if row else 0
        finally:
            conn.close()
    except Exception:
        return 0


# ═══════════════════════════════════════════════════════════════
#  Error handler
# ═══════════════════════════════════════════════════════════════

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({"error": "File too large. Maximum size is 50MB."}), 413


# ═══════════════════════════════════════════════════════════════
#  Enhanced RAG Document Store
# ═══════════════════════════════════════════════════════════════

class EnhancedDocumentStore:
    def __init__(self, persist_path=VECTOR_STORE_PATH):
        self.documents = []
        self.vectorizer = TfidfVectorizer(
            max_features=1000,
            stop_words='english',
            ngram_range=(1, 2),
            min_df=1,
            max_df=0.95
        )
        self.doc_vectors = None
        self.persist_path = persist_path
        self.metadata_path = Path(persist_path) / 'metadata.pkl'
        self.vectors_path = Path(persist_path) / 'vectors.pkl'

        # Load existing data if available
        self.load_from_disk()

    def add_document(self, text, filename, metadata=None):
        """Add a document with smart chunking"""
        chunks = self._smart_chunk_text(text, filename)

        if not chunks:
            print(f"[WARN] No chunks extracted from {filename}, skipping.")
            return 0

        for i, chunk_data in enumerate(chunks):
            doc = {
                'id': f"{filename}_{i}",
                'text': chunk_data['text'],
                'filename': filename,
                'chunk_index': i,
                'chunk_type': chunk_data['type'],
                'metadata': metadata or {},
                'timestamp': datetime.now().isoformat()
            }
            self.documents.append(doc)

        self._rebuild_vectors()
        self.save_to_disk()

        return len(chunks)

    def _smart_chunk_text(self, text, filename):
        """Smart chunking based on document structure"""
        chunks = []

        lines = text.split('\n')
        current_chunk = []
        current_type = 'paragraph'
        chunk_size = 0
        max_chunk_size = 500  # words

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect headers (simple heuristic)
            is_header = (len(line.split()) <= 10 and
                        (line.isupper() or line.endswith(':') or
                         any(line.startswith(marker) for marker in ['#', '##', '###'])))

            words = line.split()

            # Start new chunk if we hit size limit or structural boundary
            if chunk_size + len(words) > max_chunk_size or (is_header and current_chunk):
                if current_chunk:
                    chunks.append({
                        'text': ' '.join(current_chunk),
                        'type': current_type
                    })
                current_chunk = []
                chunk_size = 0

            current_chunk.extend(words)
            chunk_size += len(words)

            if is_header:
                current_type = 'header'
            else:
                current_type = 'paragraph'

        # Add final chunk
        if current_chunk:
            chunks.append({
                'text': ' '.join(current_chunk),
                'type': current_type
            })

        return chunks if chunks else ([{'text': text, 'type': 'paragraph'}] if text.strip() else [])

    def _rebuild_vectors(self):
        """Rebuild TF-IDF vectors"""
        if self.documents:
            # Prevent memory leaks by cleaning up old vectors and using a generator
            self.doc_vectors = None
            import gc
            gc.collect()
            
            texts = (doc['text'] for doc in self.documents)
            try:
                self.doc_vectors = self.vectorizer.fit_transform(texts)
                print(f"[OK] Rebuilt vectors: {self.doc_vectors.shape}")
            except ValueError as e:
                print(f"[WARN] Could not build vectors (likely too few tokens): {e}")
                self.doc_vectors = None

    def search(self, query, top_k=3, min_score=0.1):
        """Enhanced search with relevance scoring"""
        if not self.documents or self.doc_vectors is None:
            return []

        try:
            from sklearn.exceptions import NotFittedError
            try:
                query_vector = self.vectorizer.transform([query])
            except NotFittedError:
                print("[WARN] Vectorizer not fitted — rebuilding")
                self._rebuild_vectors()
                if self.doc_vectors is None:
                    return []
                query_vector = self.vectorizer.transform([query])

            # Calculate cosine similarities
            similarities = cosine_similarity(query_vector, self.doc_vectors)[0]

            # Get top results
            top_indices = np.argsort(similarities)[-top_k*2:][::-1]

            results = []

            for idx in top_indices:
                if len(results) >= top_k:
                    break

                score = similarities[idx]
                if score < min_score:
                    continue

                doc = self.documents[idx]

                # Diversify results - limit chunks per file
                file_count = sum(1 for r in results if r['document']['filename'] == doc['filename'])
                if file_count >= 2:
                    continue

                results.append({
                    'document': doc,
                    'score': float(score)
                })

            return results

        except Exception as e:
            print(f"[ERROR] Search error: {e}")
            return []

    def get_document_by_id(self, doc_id):
        """Retrieve specific document by ID"""
        for doc in self.documents:
            if doc['id'] == doc_id:
                return doc
        return None

    def get_file_chunks(self, filename):
        """Get all chunks for a specific file"""
        return [doc for doc in self.documents if doc['filename'] == filename]

    def remove_document(self, filename):
        """Remove all chunks of a document"""
        initial_count = len(self.documents)
        self.documents = [doc for doc in self.documents if doc['filename'] != filename]
        removed = initial_count - len(self.documents)

        if removed > 0:
            self._rebuild_vectors()
            self.save_to_disk()

        return removed

    def clear(self):
        """Clear all documents"""
        self.documents = []
        self.doc_vectors = None
        self.save_to_disk()

    def save_to_disk(self):
        """Persist vector store to disk"""
        try:
            with open(self.metadata_path, 'wb') as f:
                pickle.dump({
                    'documents': self.documents,
                    'vectorizer': self.vectorizer
                }, f)

            if self.doc_vectors is not None:
                with open(self.vectors_path, 'wb') as f:
                    pickle.dump(self.doc_vectors, f)
            else:
                if self.vectors_path.exists():
                    self.vectors_path.unlink()

            print(f"[SAVE] Saved {len(self.documents)} documents to disk")

        except Exception as e:
            print(f"[ERROR] Error saving to disk: {e}")

    def load_from_disk(self):
        """Load vector store from disk"""
        try:
            if self.metadata_path.exists():
                with open(self.metadata_path, 'rb') as f:
                    data = pickle.load(f)
                    self.documents = data['documents']
                    self.vectorizer = data['vectorizer']

                if self.vectors_path.exists():
                    with open(self.vectors_path, 'rb') as f:
                        self.doc_vectors = pickle.load(f)

                print(f"[OK] Loaded {len(self.documents)} documents from disk")
                return True

        except Exception as e:
            print(f"[WARN] Could not load from disk: {e}")
            self.documents = []
            self.doc_vectors = None

        return False

    def get_stats(self):
        """Get detailed statistics"""
        filenames = {}
        for doc in self.documents:
            fname = doc['filename']
            if fname not in filenames:
                filenames[fname] = {
                    'chunks': 0,
                    'upload_time': doc['metadata'].get('upload_time', 'Unknown')
                }
            filenames[fname]['chunks'] += 1

        return {
            'total_chunks': len(self.documents),
            'total_files': len(filenames),
            'files': [
                {'name': name, **info}
                for name, info in filenames.items()
            ],
            'vector_dimensions': self.doc_vectors.shape[1] if self.doc_vectors is not None else 0
        }


# Initialize document store
doc_store = EnhancedDocumentStore()


# ═══════════════════════════════════════════════════════════════
#  Text Extraction Helpers
# ═══════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_path):
    """Extract text from PDF with better error handling"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            print(f"[PDF] Processing PDF: {total_pages} pages")

            for i, page in enumerate(pdf_reader.pages[:100]):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {i+1} ---\n{page_text}"
                except Exception as e:
                    print(f"[WARN] Error on page {i+1}: {e}")

    except Exception as e:
        print(f"[ERROR] PDF extraction error: {e}")

    return text


def extract_text_from_docx(file_path):
    """Extract text from DOCX"""
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text:
                    tables_text.append(row_text)

        full_text = '\n'.join(paragraphs)
        if tables_text:
            full_text += '\n\n--- Tables ---\n' + '\n'.join(tables_text)

        return full_text
    except Exception as e:
        print(f"[ERROR] DOCX extraction error: {e}")
        return ""


def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        print(f"[ERROR] TXT extraction error: {e}")
        return ""


def extract_text_from_html(file_path):
    """Extract text from HTML file, stripping tags"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            html_content = file.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text and clean up whitespace
        text = soup.get_text(separator=' ', strip=True)
        text = re.sub(r'\s+', ' ', text).strip()

        return text
    except Exception as e:
        print(f"[ERROR] HTML extraction error: {e}")
        return ""


def extract_text_from_zip(file_path):
    """Extract text from ZIP file containing multiple documents"""
    extracted_texts = []

    try:
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            file_list = zip_ref.namelist()
            print(f"[ZIP] Found {len(file_list)} files in archive")

            with tempfile.TemporaryDirectory() as temp_dir:
                temp_dir_resolved = os.path.realpath(temp_dir)
                for member in file_list:
                    member_path = os.path.realpath(os.path.join(temp_dir, member))
                    # Zip Slip protection — wrapped in try/except for Windows drive mismatch
                    try:
                        if os.path.commonpath([member_path, temp_dir_resolved]) != temp_dir_resolved:
                            print(f"[SECURITY] Blocked path traversal attempt: {member}")
                            continue
                    except ValueError:
                        # os.path.commonpath raises ValueError on Windows if paths are on different drives
                        print(f"[SECURITY] Skipping member on different drive: {member}")
                        continue
                    zip_ref.extract(member, temp_dir)

                # Process each extracted file
                for member in file_list:
                    member_path = os.path.join(temp_dir, member)

                    if os.path.isdir(member_path):
                        continue

                    if not os.path.exists(member_path):
                        continue

                    ext = member.rsplit('.', 1)[-1].lower() if '.' in member else ''

                    print(f"[ZIP] Processing: {member}")

                    try:
                        if ext == 'pdf':
                            text = extract_text_from_pdf(member_path)
                        elif ext == 'docx':
                            text = extract_text_from_docx(member_path)
                        elif ext in ['html', 'htm']:
                            text = extract_text_from_html(member_path)
                        elif ext in ['txt', 'md', 'csv']:
                            text = extract_text_from_txt(member_path)
                        elif ext == 'json':
                            with open(member_path, 'r', encoding='utf-8', errors='ignore') as f:
                                text = f.read()
                        else:
                            text = ""

                        if text.strip():
                            extracted_texts.append(f"--- File: {member} ---\n{text}")

                    except Exception as e:
                        print(f"[ERROR] Error processing {member}: {e}")

        result = "\n\n".join(extracted_texts)
        print(f"[ZIP] Extracted text from {len(extracted_texts)} files")
        return result

    except Exception as e:
        print(f"[ERROR] ZIP extraction error: {e}")
        return ""


def allowed_file(filename):
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].strip().lower()
    if not ext:
        return False
    if ALLOWED_EXTENSIONS is None:
        return True
    return ext in ALLOWED_EXTENSIONS


# ═══════════════════════════════════════════════════════════════
#  Ollama Helpers
# ═══════════════════════════════════════════════════════════════

def check_available_models():
    """Check available models from Ollama."""
    for path in ["/v1/models", "/api/tags"]:
        try:
            url = f"{OLLAMA_BASE_URL}{path}"
            response = requests.get(url, timeout=5)
            if response.status_code == 200:
                data = response.json()

                if isinstance(data, dict) and 'models' in data:
                    # Ollama /api/tags format: {"models": [{"name": "..."}]}
                    models = data.get('models') or []
                    if not isinstance(models, list):
                        models = []
                elif isinstance(data, dict) and 'data' in data:
                    # OpenAI-compatible /v1/models format: {"data": [{"id": "..."}]}
                    models = data.get('data') or []
                    if not isinstance(models, list):
                        models = []
                elif isinstance(data, list):
                    models = data
                else:
                    models = []

                model_names = []
                for model in models:
                    if isinstance(model, dict) and 'name' in model:
                        model_names.append(model['name'])
                    elif isinstance(model, dict) and 'id' in model:
                        model_names.append(model['id'])
                    elif isinstance(model, str):
                        model_names.append(model)

                print(f"[OK] Available models ({path}): {model_names}")
                return model_names
            elif response.status_code == 404:
                continue
        except requests.exceptions.RequestException as e:
            print(f"[WARN] Could not read models from {path}: {e}")

    print(f"[ERROR] No available model endpoint found on {OLLAMA_BASE_URL}")
    return []


def is_ollama_reachable():
    """Quick check if Ollama is running."""
    try:
        resp = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=3)
        return resp.status_code == 200
    except Exception:
        return False


def is_vision_model(model_name):
    """Check if the model supports vision."""
    lower = model_name.lower()
    return any(kw in lower for kw in ["llava", "vision", "bakllava", "qwen2.5vl", "qwen2.5-vl", "minicpm", "moondream"])


def parse_ollama_response(result):
    """Parse the response from Ollama API"""
    if isinstance(result, dict):
        return result.get('response', str(result))
    return str(result)


def get_local_ip():
    """Get local IP address"""
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        return local_ip
    except Exception:
        return "localhost"
    finally:
        s.close()


# Add message to in-memory conversation history
def add_to_history(user_msg, assistant_msg):
    """Add a message pair to in-memory conversation history"""
    with history_lock:
        conversation_history.append({"role": "user", "content": user_msg})
        conversation_history.append({"role": "assistant", "content": assistant_msg})
        if len(conversation_history) > MAX_HISTORY_MESSAGES * 2:
            conversation_history[:] = conversation_history[-MAX_HISTORY_MESSAGES * 2:]


# ═══════════════════════════════════════════════════════════════
#  Smart Title Generation (Background Thread)
# ═══════════════════════════════════════════════════════════════

def generate_title_async(conversation_id, user_message, model_used=None):
    """Generate a conversation title in the background using Ollama.

    `model_used` should be the model that just generated the actual reply
    (e.g. deepseek-r1:8b in Think mode). Reusing it instead of always
    falling back to MODEL_NAME avoids forcing Ollama to evict the model that
    is already hot in VRAM just to load a different one for a 5-word title.
    """
    try:
        title_model = model_used or MODEL_NAME
        prompt = (
            f"In 5 words or fewer, write a concise title for a conversation "
            f"that starts with this message: '{user_message[:200]}'. "
            f"Reply with only the title. No quotes, no punctuation at the end."
        )
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/chat",
            json={
                "model": title_model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "options": {"num_predict": 15, "temperature": 0.3, "num_gpu": 99},
                "keep_alive": GPU_KEEP_ALIVE
            },
            timeout=30
        )
        if resp.status_code == 200:
            data = resp.json()
            title = data.get("message", {}).get("content", "").strip()
            # Clean up the title
            title = title.strip('"\'').strip()
            if not title or len(title) < 2:
                title = user_message[:50].strip()
            elif len(title) > 60:
                title = title[:57] + "..."

            with db_lock:
                conn = get_db()
                try:
                    conn.execute(
                        "UPDATE conversations SET title = ? WHERE id = ?",
                        (title, conversation_id)
                    )
                    conn.commit()
                finally:
                    conn.close()
            print(f"[TITLE] Generated title for {conversation_id[:8]}...: {title}")
    except Exception as e:
        print(f"[WARN] Title generation failed: {e}")
        # Fallback: use first 50 chars of user message
        try:
            fallback = user_message[:50].strip()
            with db_lock:
                conn = get_db()
                try:
                    conn.execute(
                        "UPDATE conversations SET title = ? WHERE id = ? AND title = 'New Chat'",
                        (fallback, conversation_id)
                    )
                    conn.commit()
                finally:
                    conn.close()
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════
#  Web Search (DuckDuckGo)
# ═══════════════════════════════════════════════════════════════

def search_duckduckgo(query, max_results=5):
    """Fetch search results from DuckDuckGo HTML."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        resp = requests.get(
            f"https://duckduckgo.com/html/?q={quote_plus(query)}",
            headers=headers,
            timeout=10
        )
        if resp.status_code != 200:
            return []

        soup = BeautifulSoup(resp.text, 'html.parser')
        results = []

        for result_div in soup.find_all('div', class_='result', limit=max_results * 2):
            snippet_el = result_div.find('a', class_='result__snippet') or result_div.find('td', class_='result__snippet')
            url_el = result_div.find('a', class_='result__url') or result_div.find('span', class_='result__url')

            snippet = snippet_el.get_text(strip=True) if snippet_el else ""
            url = url_el.get_text(strip=True) if url_el else ""

            if snippet:
                results.append({"url": url, "snippet": snippet})
                if len(results) >= max_results:
                    break

        return results
    except Exception as e:
        print(f"[WARN] DuckDuckGo search failed: {e}")
        return []


def execute_step_list(steps_text):
    """
    Parse and execute a rigid STEP-BY-STEP format outputted by the LLM.
    Format: STEP N: ACTION_TYPE | argument
    Yields status updates for the frontend.
    """
    import time
    import subprocess
    import sys
    import re
    try:
        import desktop_actions
    except ImportError:
        yield "❌ desktop_actions module not found.\n"
        return

    lines = steps_text.split('\n')
    steps_found = 0
    success_count = 0
    
    yield f"\n[AUTOMATE] Parsing execution steps...\n"
    
    for line in lines:
        line = line.strip()
        # Look for "STEP X: ACTION | arg"
        match = re.match(r'^STEP\s+\d+:\s+([A-Z_]+)\s*\|\s*(.*)$', line)
        if not match:
            continue
            
        action = match.group(1).strip()
        arg = match.group(2).strip()
        steps_found += 1
        
        yield f"Step {steps_found}: {action} '{arg}' — "
        
        try:
            if action == "LAUNCH_APP":
                res = desktop_actions.open_app(arg)
                yield f"{res}\n"
                if res == "success": success_count += 1
            elif action == "TYPE_TEXT":
                res = desktop_actions.type_text(arg)
                yield f"{res}\n"
                if res == "success": success_count += 1
            elif action == "PRESS_KEY":
                res = desktop_actions.press_key(arg)
                yield f"{res}\n"
                if res == "success": success_count += 1
            elif action == "OPEN_URL":
                res = desktop_actions.open_url(arg)
                yield f"{res}\n"
                if res == "success": success_count += 1
            elif action == "WAIT_SECONDS":
                time.sleep(float(arg))
                yield "success\n"
                success_count += 1
            elif action == "RUN_PYTHON":
                # Execute a one-liner python script
                res = subprocess.run([sys.executable, "-c", arg], capture_output=True, text=True)
                if res.returncode == 0:
                    yield "success\n"
                    success_count += 1
                else:
                    yield f"error: {res.stderr}\n"
            else:
                yield f"unknown action\n"
        except Exception as e:
            yield f"error: {str(e)}\n"
            
    if steps_found == 0:
        yield "\n❌ No valid steps found in LLM response. The LLM failed to follow the rigid format.\n"
    else:
        yield f"\n✅ Completed {success_count}/{steps_found} steps successfully.\n"

# ═══════════════════════════════════════════════════════════════
#  Automation, Project Indexing & Harness Tools
# ═══════════════════════════════════════════════════════════════

def execute_automation_script(script_code):
    """
    Execute a pyautogui automation script in the workspace sandbox.
    Returns (success: bool, output: str, error: str).
    """
    import subprocess
    import sys
    import re

    # Try to extract code from markdown blocks if the LLM used them
    match = re.search(r'```(?:python)?\s*(.*?)\s*```', script_code, re.DOTALL | re.IGNORECASE)
    if match:
        script_code = match.group(1)

    # Aggressive cleanup of common LLM english phrases that cause SyntaxErrors
    clean_lines = []
    for line in script_code.split('\n'):
        stripped = line.strip().lower()
        if not stripped or stripped.startswith("#"):
            clean_lines.append(line)
            continue
            
        # Common english garbage from small LLMs
        if (stripped.startswith("here is") or 
            stripped.startswith("sure") or 
            stripped.startswith("code:") or
            stripped.startswith("please let me") or
            stripped.startswith("best regards") or
            stripped.startswith("[your name]") or
            stripped.startswith("you can assume") or
            stripped.startswith("get the path") or
            stripped.startswith("if it's not") or
            stripped.startswith("wait a bit") or
            stripped.startswith("note:") or
            stripped.startswith("to open") or
            stripped.startswith("first,") or
            stripped.startswith("then,") or
            stripped.startswith("next,") or
            stripped.startswith("finally,") or
            stripped.startswith("this script") or
            stripped.startswith("the above") or
            stripped.startswith("you should") or
            stripped.startswith("you can") or
            stripped.startswith("you need") or
            stripped.startswith("make sure") or
            stripped.startswith("don't forget") or
            stripped.startswith("remember to") or
            stripped.startswith("the script") or
            stripped.startswith("this will") or
            stripped.startswith("this code") or
            stripped.startswith("as a result") or
            stripped.startswith("in order to")):
            clean_lines.append("# LLM TEXT: " + line) # Comment it out safely
            continue
            
        # English sentences ending with colon that aren't python blocks
        if stripped.endswith(":") and " " in stripped and not stripped.startswith(("def ", "class ", "if ", "elif ", "else:", "for ", "while ", "with ", "try:", "except")):
            clean_lines.append("# LLM TEXT: " + line)
            continue

        # General heuristic: If it doesn't start with a valid Python identifier char/keyword and has a space
        first_char = stripped[0]
        if not (first_char.isalpha() or first_char in '_#0123456789"\' ') and " " in stripped:
            clean_lines.append("# LLM TEXT: " + line)
            continue
            
        clean_lines.append(line)
        
    script_code = "\n".join(clean_lines)

    script_filename = f"temp_run_{uuid.uuid4().hex[:8]}.py"
    script_path = os.path.join(WORKSPACE_DIR, script_filename)

    # Safety preamble injected before every automation script
    safety_preamble = """
import time
import subprocess
import os
import sys
import webbrowser
from pathlib import Path

try:
    import pyautogui
    pyautogui.FAILSAFE = True
    pyautogui.PAUSE = 0.5
except ImportError:
    pyautogui = None

"""
    full_script = safety_preamble + script_code

    # Attempt compilation
    try:
        compile(full_script, script_filename, 'exec')
    except SyntaxError as e:
        # Second pass: remove lines > 4 words with no Python operators
        clean_lines_pass2 = []
        for line in full_script.split('\n'):
            stripped = line.strip()
            if not stripped or stripped.startswith("#"):
                clean_lines_pass2.append(line)
                continue
            
            words = stripped.split()
            has_operator = any(op in stripped for op in ["=", "(", ")", ":", "[", "]", ".", "+", "-", "*", "/", "%", "<", ">", "!"])
            if len(words) > 4 and not has_operator:
                clean_lines_pass2.append("# LLM TEXT (Pass 2): " + line)
            else:
                clean_lines_pass2.append(line)
                
        full_script = "\n".join(clean_lines_pass2)
        
        try:
            compile(full_script, script_filename, 'exec')
        except SyntaxError as e2:
            return False, "", f"SyntaxError in LLM generated script: {str(e2)}"

    try:
        with open(script_path, 'w', encoding='utf-8') as f:
            f.write(full_script)

        result = subprocess.run(
            [sys.executable, script_filename],
            capture_output=True,
            text=True,
            timeout=60,
            cwd=WORKSPACE_DIR
        )
        output = result.stdout.strip()
        error = result.stderr.strip()
        success = result.returncode == 0
        return success, output, error

    except subprocess.TimeoutExpired:
        return False, "", "Automation script timed out after 60 seconds."
    except Exception as e:
        return False, "", f"Failed to execute automation script: {str(e)}"
    finally:
        if os.path.exists(script_path):
            try:
                os.remove(script_path)
            except Exception:
                pass


# ═══════════════════════════════════════════════════════════════
#  Quick-Script Helpers (fallback when OpenClaw is offline)
# ═══════════════════════════════════════════════════════════════

# Known browser executables for "open X in Y" patterns
_BROWSER_MAP = {
    'chrome': 'chrome', 'google chrome': 'chrome',
    'firefox': 'firefox', 'mozilla': 'firefox',
    'edge': 'msedge', 'microsoft edge': 'msedge',
    'opera': 'opera', 'opera gx': 'opera',
    'brave': 'brave',
}

# Common websites → URLs
_SITE_MAP = {
    'youtube': 'https://www.youtube.com',
    'google': 'https://www.google.com',
    'gmail': 'https://mail.google.com',
    'github': 'https://github.com',
    'twitter': 'https://twitter.com', 'x': 'https://twitter.com',
    'reddit': 'https://www.reddit.com',
    'instagram': 'https://www.instagram.com',
    'facebook': 'https://www.facebook.com',
    'whatsapp': 'https://web.whatsapp.com',
    'spotify': 'https://open.spotify.com',
    'spotify liked': 'https://open.spotify.com/collection/tracks',
    'spotify liked playlist': 'https://open.spotify.com/collection/tracks',
    'liked songs': 'https://open.spotify.com/collection/tracks',
    'chatgpt': 'https://chat.openai.com',
    'linkedin': 'https://www.linkedin.com',
    'netflix': 'https://www.netflix.com',
    'twitch': 'https://www.twitch.tv',
}


_APP_LAUNCH_MAP = {
    "visual studio code": "code",
    "vscode": "code",
    "vs code": "code",
    "notepad": "notepad",
    "notepad++": "notepad++",
    "chrome": "chrome",
    "google chrome": "chrome",
    "firefox": "firefox",
    "discord": "discord",
    "spotify": "spotify",
    "calculator": "calc",
    "paint": "mspaint",
    "file explorer": "explorer",
    "task manager": "taskmgr",
    "cmd": "cmd",
    "command prompt": "cmd",
    "powershell": "powershell",
    "terminal": "wt",
    "windows terminal": "wt",
    "word": "winword",
    "excel": "excel",
    "powerpoint": "powerpnt",
    "outlook": "outlook",
    "teams": "teams",
    "obs": "obs64",
    "vlc": "vlc",
    "steam": "steam"
}


def _build_quick_script(task_description):
    """Build a quick pyautogui script for simple tasks. Returns script string or None."""
    task_lower = task_description.lower().strip()
    action_words = {"and", "then", "type", "write", "create", "make"}
    words = task_lower.split()
    
    # Return None for any complex/multi-step task
    if any(w in action_words for w in words):
        return None

    # Handle explicit "open [app]" or "open [site]"
    if task_lower.startswith("open "):
        target = task_lower[5:].strip()
        
        # Check known sites
        if target in _SITE_MAP:
            url = _SITE_MAP[target]
            return f"import webbrowser\nwebbrowser.open({url!r})\n"
        
        # Check known apps
        if target in _APP_LAUNCH_MAP:
            exe = _APP_LAUNCH_MAP[target]
            return f"import subprocess\nsubprocess.Popen([{exe!r}])\n"
            
        # Check explicit URLs
        if '.' in target or target.startswith(('http://', 'https://')):
            url = target if target.startswith(('http://', 'https://')) else f'https://{target}'
            return f"import webbrowser\nwebbrowser.open({url!r})\n"
            
        # If it's a short target (1-3 words), fall back to Windows search
        target_words = target.split()
        if len(target_words) <= 3:
            return (
                f"import pyautogui, time\n"
                f"pyautogui.press('win')\n"
                f"time.sleep(1)\n"
                f"pyautogui.write({target!r}, interval=0.05)\n"
                f"time.sleep(0.8)\n"
                f"pyautogui.press('enter')\n"
            )
            
    return None


def _try_quick_script_fallback(task_description):
    """Try executing a quick script as fallback. Returns formatted result string or None."""
    script = _build_quick_script(task_description)
    if not script:
        return None

    success, out, err = execute_automation_script(script)
    log_automation(task_description, f"fallback: {script[:100]}", out if success else err)

    if success:
        if "pyautogui.press('win')" in script:
            return f"**⚠️ Ran a best-effort Windows Search fallback** (exit code 0, but this cannot verify whether the task was actually completed)\n```\n{out or 'Done!'}\n```"
        else:
            return f"**✅ Fallback automation succeeded**\n```\n{out or 'Done!'}\n```"
    else:
        return f"**❌ Fallback automation failed**\n```\n{err}\n```"


# ═══════════════════════════════════════════════════════════════
#  OpenClaw Integration — Smart Autonomous Agent Automation
# ═══════════════════════════════════════════════════════════════

OPENCLAW_SESSION_KEY = "agent:main:marceline-automation"

def ensure_openclaw_computer_skill():
    skill_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "workspace", "skills", "computer_control")
    skill_file = os.path.join(skill_dir, "SKILL.md")
    desktop_actions_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "desktop_actions.py").replace("\\", "/")
    
    needs_update = True
    if os.path.exists(skill_file):
        try:
            with open(skill_file, "r", encoding="utf-8") as f:
                content = f.read()
                import re
                match = re.search(r'([a-zA-Z0-9_/\-:]+desktop_actions\.py)', content)
                if match:
                    found_path = match.group(1)
                    if found_path == desktop_actions_path:
                        needs_update = False
        except Exception:
            pass

    if needs_update:
        os.makedirs(skill_dir, exist_ok=True)
        with open(skill_file, "w", encoding="utf-8") as f:
            f.write(f"""---
name: computer_control
description: Interacts with the user's desktop and browser using pyautogui for OS control and Puppeteer for browser automation.
---

# Computer Control — OS and Browser Automation

You have access to BOTH real desktop control AND full browser automation.

## Desktop Control (OS-level)
Use the `exec` tool to run `desktop_actions.py`:

The python script is located at: `{desktop_actions_path}`

Supported verbs:
- `open_app <app_name>`: Opens an application (e.g., spotify, vscode, notepad).
- `open_url <url>`: Opens a URL in the default browser.
- `click <x> <y>`: Clicks at pixel coordinates.
- `type_text "<text>"`: Types text using keyboard.
- `press_key <key>`: Presses a keyboard key (enter, space, esc, tab, etc.).
- `screenshot`: Takes a screenshot of the current screen.

Example to open VS Code and create a Python file:
```bash
python {desktop_actions_path} open_app vscode
```

## Browser Automation (Puppeteer)
You also have access to Puppeteer for full browser control. Use the `browser` tool (provided by the OpenClaw gateway) with these actions:

- `browser.goto(url)` — Navigate to a URL
- `browser.click(selector)` — Click a CSS selector or element
- `browser.type(selector, text)` — Type into an input field
- `browser.screenshot()` — Take a screenshot of the current browser page
- `browser.evaluate(js)` — Run JavaScript in the page
- `browser.waitForSelector(selector)` — Wait for an element to appear
- `browser.content()` — Get the full HTML of the current page

Example to open a GitHub repo:
```javascript
await browser.goto('https://github.com');
await browser.type('input[name="q"]', 'marceline project');
await browser.press('Enter');
```

## When to use which:
- Use Puppeteer/browser tools for any task involving websites, web apps, forms, or URLs.
- Use desktop_actions.py for any task involving native apps (VS Code, Spotify, Notepad, etc.), file operations, or keyboard/mouse control.
- For tasks like "open VS Code and create a Python file": use open_app first, then type_text/press_key.

Always use the `exec` tool to run desktop_actions.py commands. Always use the `browser` tool for Puppeteer commands.
""")
        print("[SETUP] Created OpenClaw computer_control skill.")

def run_vision_guided_automation(task_description, max_steps=6):
    import json
    import re
    import time
    try:
        import pyautogui
    except ImportError:
        yield "❌ pyautogui is not installed. Vision automation requires it.\n"
        return

    # Add the local directory to sys.path if not there so we can import desktop_actions
    script_dir = os.path.dirname(os.path.abspath(__file__))
    if script_dir not in sys.path:
        sys.path.append(script_dir)
        
    try:
        import desktop_actions
    except ImportError:
        yield "❌ desktop_actions module not found.\n"
        return
        
    yield f"\n[VISION] Starting vision-guided automation for: {task_description}\n"
    
    history = []
    vision_model_to_use = "llava:latest"
    model_checked = False
    available_vision_models_cache = None

    for step in range(1, max_steps + 1):
        yield f"\n[VISION] Step {step}/{max_steps}...\n"
        
        # 1. Take screenshot
        b64_img = desktop_actions.screenshot(return_base64=True)
        if b64_img == "error":
            yield f"Step {step}/{max_steps}: Failed to take screenshot.\n"
            return
            
        # Dynamically determine the best available vision model (only once per call)
        if available_vision_models_cache is None:
            try:
                tags_resp = requests.get("http://localhost:11434/api/tags", timeout=2)
                if tags_resp.status_code == 200:
                    available_vision_models_cache = [m["name"] for m in tags_resp.json().get("models", [])]
                else:
                    available_vision_models_cache = []
            except Exception:
                available_vision_models_cache = []

        available = available_vision_models_cache
        qwen_models = [m for m in available if "qwen2.5vl" in m or "qwen2.5-vl" in m]
        if qwen_models:
            vision_model_to_use = qwen_models[0]
        elif "llava:latest" in available:
            vision_model_to_use = "llava:latest"
        else:
            for m in available:
                if "llava" in m or "qwen" in m:
                    vision_model_to_use = m
                    break
            
        if not model_checked:
            try:
                show_resp = requests.post("http://localhost:11434/api/show", json={"model": vision_model_to_use}, timeout=5)
                if show_resp.status_code == 200:
                    caps = show_resp.json().get("capabilities", []) or []
                    if "vision" not in caps:
                        yield f"❌ Model '{vision_model_to_use}' does not report vision/multimodal support. Please pull a vision model.\n"
                        return
            except Exception:
                pass
            model_checked = True
            
        yield f"Step {step}/{max_steps}: Analyzing screen (using {vision_model_to_use})...\n"
        
        if b64_img.startswith("data:image"):
            b64_img = b64_img.split(",", 1)[-1]
            
        # 2. Ask Vision Model
        prompt = f"""You are a desktop automation agent controlling a computer.
Task: {task_description}

Previous actions: {json.dumps(history)}

Analyze the provided screenshot.
Is the target application open and visible? If not, what should be done to open it?
If it is visible, what is the next specific action to take to progress the task?
If the task is complete, return 'done'.
If you cannot proceed, return 'give_up'.

You MUST return your response as a valid JSON object matching this schema exactly, and nothing else:
{{
  "app_visible": boolean,
  "target_element_found": boolean,
  "coordinates": [x, y] or null,
  "next_action": "click" | "type" | "open_app" | "wait" | "done" | "give_up",
  "action_args": "argument for the action (e.g. app name, text to type) or null",
  "reasoning": "short explanation of your choice"
}}
"""
        try:
            resp = requests.post("http://localhost:11434/api/generate", json={
                "model": vision_model_to_use,
                "prompt": prompt,
                "images": [b64_img],
                "stream": False,
                "format": "json",
                "options": {"num_gpu": 99, "num_predict": 512, "temperature": 0.2},
                "keep_alive": GPU_KEEP_ALIVE
            }, timeout=120)
            
            if resp.status_code != 200:
                err_msg = ""
                try:
                    err_msg = resp.json().get("error", resp.text)
                except Exception:
                    err_msg = resp.text
                full_error = f"Vision model error {resp.status_code}: {err_msg}"
                yield f"Step {step}/{max_steps}: {full_error}\n"
                log_automation(task_description, "vision_loop", f"error: {full_error}")
                break
                
            result_json = resp.json().get("response", "")
            
            # Robust JSON extraction
            match = re.search(r'\{.*\}', result_json, re.DOTALL)
            if match:
                parsed = json.loads(match.group(0))
            else:
                parsed = json.loads(result_json)
                
            reasoning = parsed.get("reasoning", "")
            next_action = parsed.get("next_action", "")
            action_args = parsed.get("action_args", "")
            coords = parsed.get("coordinates", None)
            
            yield f"Step {step}/{max_steps} [Vision]: {reasoning}\n"
            
            if next_action == "done":
                yield "\n✅ Vision-guided automation completed successfully.\n"
                log_automation(task_description, "vision_loop", "success")
                return
            elif next_action == "give_up":
                yield "\n❌ Vision-guided automation gave up.\n"
                log_automation(task_description, "vision_loop", "gave_up")
                return
            elif next_action == "click":
                if coords and len(coords) == 2:
                    sw, sh = pyautogui.size()
                    x, y = coords[0], coords[1]
                    if 0 <= x <= sw and 0 <= y <= sh:
                        desktop_actions.click(x, y)
                        history.append({"step": step, "action": "click", "coords": [x, y], "reasoning": reasoning})
                        yield f"Step {step}/{max_steps}: Clicked at ({x}, {y}).\n"
                    else:
                        history.append({"step": step, "error": f"Invalid coordinates ({x}, {y}) out of bounds ({sw}, {sh})"})
                        yield f"Step {step}/{max_steps}: Invalid coordinates returned, retrying...\n"
                else:
                    yield f"Step {step}/{max_steps}: No coordinates provided for click.\n"
            elif next_action == "type":
                desktop_actions.type_text(str(action_args))
                history.append({"step": step, "action": "type", "args": action_args, "reasoning": reasoning})
                yield f"Step {step}/{max_steps}: Typed text.\n"
            elif next_action == "open_app":
                desktop_actions.open_app(str(action_args))
                history.append({"step": step, "action": "open_app", "args": action_args, "reasoning": reasoning})
                yield f"Step {step}/{max_steps}: Opened app {action_args}.\n"
                time.sleep(2) # Wait for app to open
            elif next_action == "wait":
                time.sleep(2)
                history.append({"step": step, "action": "wait", "reasoning": reasoning})
                yield f"Step {step}/{max_steps}: Waiting...\n"
            else:
                yield f"Step {step}/{max_steps}: Unknown action {next_action}.\n"
                
            time.sleep(1) # Pause before next screenshot
            
        except Exception as e:
            yield f"Step {step}/{max_steps}: Error in vision loop: {str(e)}\n"
            log_automation(task_description, "vision_loop", f"error: {str(e)}")
            break
            
    yield "\n❌ Vision-guided automation reached max steps.\n"
    log_automation(task_description, "vision_loop", "max_steps_reached")

def execute_native_agent(task_description, timeout=120):
    """
    Native Marceline agent loop. Uses the local LLM (via Ollama) to plan
    a sequence of desktop/browser actions, then executes them directly using
    desktop_actions.py and Puppeteer (via Node.js). No OpenClaw required.
    
    Returns (success: bool, steps_log: list, final_message: str, error: str)
    """
    import subprocess as _sp
    import sys
    import re
    import time

    script_dir = os.path.dirname(os.path.abspath(__file__))
    desktop_actions_path = os.path.join(script_dir, "desktop_actions.py")

    # ── Step 1: Ask Ollama to produce an action plan ─────────────────────────
    planning_prompt = (
        f"You are a desktop automation engine. The user wants to: {task_description}\n\n"
        "Produce a numbered action plan. Each line must be EXACTLY this format:\n"
        "STEP N: ACTION | argument\n\n"
        "Available actions:\n"
        "  OPEN_APP     | app_name          (opens a native app: vscode, notepad, spotify, calculator, chrome, etc.)\n"
        "  OPEN_URL     | https://...       (opens a URL in the browser)\n"
        "  CLICK        | x,y               (clicks at screen coordinates)\n"
        "  TYPE_TEXT    | your text here    (types text at current cursor)\n"
        "  PRESS_KEY    | key_name          (presses a key: enter, tab, ctrl+s, ctrl+n, etc.)\n"
        "  WAIT         | seconds           (waits N seconds)\n"
        "  RUN_PYTHON   | python code here  (runs one line of Python)\n"
        "  SYSTEM_CMD   | command           (runs a Windows system command: shutdown /r /t 0, shutdown /s /t 0, logoff, etc.)\n"
        "  BROWSER_GOTO | https://...       (navigate browser to URL via Puppeteer)\n"
        "  BROWSER_CLICK| css_selector      (click an element in the browser)\n"
        "  BROWSER_TYPE | selector###text   (type into browser element, use ### to separate selector from text)\n"
        "  SCREENSHOT   |                   (take a screenshot)\n\n"
        "Rules:\n"
        "- No explanations. No blank lines. No markdown. Only STEP lines.\n"
        "- NEVER use CLICK with coordinates 0,0 or any screen corner — those trigger the safety abort.\n"
        "- Only use CLICK if you have a specific reason to know the exact coordinates. Prefer TYPE_TEXT and PRESS_KEY instead.\n"
        "- For keyboard shortcuts like Ctrl+S: use PRESS_KEY | ctrl+s\n"
        "- For creating a new file in VS Code: PRESS_KEY | ctrl+n then TYPE_TEXT | code, then PRESS_KEY | ctrl+s\n"
        "- Always add WAIT steps after opening apps (give them time to load).\n\n"
        f"Example for 'open notepad and type hello world':\n"
        "STEP 1: OPEN_APP | notepad\n"
        "STEP 2: WAIT | 1.5\n"
        "STEP 3: TYPE_TEXT | hello world\n\n"
        "For 'restart my laptop': use SYSTEM_CMD | shutdown /r /t 0\n"
        "For 'shutdown my laptop': use SYSTEM_CMD | shutdown /s /t 0\n"
        "For 'lock my computer': use SYSTEM_CMD | rundll32.exe user32.dll,LockWorkStation\n\n"
        f"Now produce the plan for: {task_description}"
    )

    try:
        plan_resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={
                "model": MODEL_NAME,
                "prompt": planning_prompt,
                "stream": False,
                "options": {"temperature": 0.1, "num_predict": 512, "num_gpu": 99},
                "keep_alive": GPU_KEEP_ALIVE
            },
            timeout=60
        )
        if plan_resp.status_code != 200:
            return False, [], "", f"LLM planning failed: HTTP {plan_resp.status_code}"
        
        plan_text = plan_resp.json().get("response", "").strip()
    except Exception as e:
        return False, [], "", f"LLM planning request failed: {str(e)}"

    # ── Step 2: Parse the plan into steps ────────────────────────────────────
    step_pattern = re.compile(
        r"STEP\s+\d+\s*:\s*(OPEN_APP|OPEN_URL|CLICK|TYPE_TEXT|PRESS_KEY|WAIT|RUN_PYTHON|SYSTEM_CMD|BROWSER_GOTO|BROWSER_CLICK|BROWSER_TYPE|SCREENSHOT)\s*\|\s*(.*)",
        re.IGNORECASE
    )

    steps = []
    for line in plan_text.splitlines():
        m = step_pattern.match(line.strip())
        if m:
            steps.append((m.group(1).upper().strip(), m.group(2).strip()))

    if not steps:
        return False, [], plan_text, "LLM did not produce a valid action plan. Raw output:\n" + plan_text[:500]

    # ── Step 3: Execute each step ─────────────────────────────────────────────
    steps_log = []
    browser_node_script_lines = []  # accumulate browser steps for batched Puppeteer run
    
    # Only run Puppeteer for steps that genuinely need a browser
    # Filter out cases where the LLM mistakenly routed OS tasks to BROWSER_GOTO
    filtered_browser_steps = [
        (a, arg) for a, arg in steps
        if a.startswith("BROWSER_") and (
            arg.startswith("http://") or arg.startswith("https://")
        ) and "restart" not in arg.lower()
        and "shutdown" not in arg.lower()
        and "reboot" not in arg.lower()
    ]
    has_browser_steps = len(filtered_browser_steps) > 0

    browser_script_path = None

    # Pre-generate Puppeteer script if needed
    if has_browser_steps:
        browser_steps_for_script = filtered_browser_steps
        pup_lines = [
            "const puppeteer = require('puppeteer');",
            "(async () => {",
            "  const browser = await puppeteer.launch({headless: false, defaultViewport: null});",
            "  const page = await browser.newPage();",
        ]
        for action, arg in browser_steps_for_script:
            if action == "BROWSER_GOTO":
                pup_lines.append(f"  await page.goto({repr(arg)}, {{waitUntil: 'domcontentloaded'}});")
            elif action == "BROWSER_CLICK":
                pup_lines.append(f"  await page.waitForSelector({repr(arg)}, {{timeout: 10000}});")
                pup_lines.append(f"  await page.click({repr(arg)});")
            elif action == "BROWSER_TYPE":
                if "###" in arg:
                    sel, text = arg.split("###", 1)
                    pup_lines.append(f"  await page.waitForSelector({repr(sel.strip())}, {{timeout: 10000}});")
                    pup_lines.append(f"  await page.type({repr(sel.strip())}, {repr(text.strip())});")
                else:
                    pup_lines.append(f"  // Could not parse BROWSER_TYPE arg: {arg}")
        pup_lines.append("  // Keep browser open for 5s so user can see result")
        pup_lines.append("  await new Promise(r => setTimeout(r, 5000));")
        pup_lines.append("  await browser.close();")
        pup_lines.append("})();")
        browser_script_path = os.path.join(script_dir, "workspace", f"browser_task_{uuid.uuid4().hex[:8]}.js")

    for action, arg in steps:
        log_entry = {"action": action, "arg": arg, "result": "", "success": True}

        try:
            if action == "OPEN_APP":
                import desktop_actions
                result = desktop_actions.open_app(arg)
                log_entry["result"] = result
                log_entry["success"] = not str(result).startswith("error")

            elif action == "OPEN_URL":
                import desktop_actions
                result = desktop_actions.open_url(arg)
                log_entry["result"] = result
                log_entry["success"] = not str(result).startswith("error")

            elif action == "CLICK":
                parts = arg.replace(" ", "").split(",")
                if len(parts) == 2:
                    import desktop_actions
                    result = desktop_actions.click(parts[0], parts[1])
                    log_entry["result"] = result
                    log_entry["success"] = not str(result).startswith("error")
                else:
                    log_entry["result"] = "error: bad coordinates format — expected x,y"
                    log_entry["success"] = False

            elif action == "TYPE_TEXT":
                import desktop_actions
                result = desktop_actions.type_text(arg)
                log_entry["result"] = result
                log_entry["success"] = not str(result).startswith("error")

            elif action == "PRESS_KEY":
                import desktop_actions
                result = desktop_actions.press_key(arg)
                log_entry["result"] = result
                log_entry["success"] = not str(result).startswith("error")

            elif action == "WAIT":
                try:
                    secs = float(arg)
                    time.sleep(secs)
                    log_entry["result"] = f"waited {secs}s"
                except ValueError:
                    time.sleep(1)
                    log_entry["result"] = "waited 1s (bad arg)"

            elif action == "SYSTEM_CMD":
                import subprocess as _sc
                # Commands that affect the session itself — don't capture, don't wait long
                session_commands = ("shutdown", "logoff", "restart", "reboot", "lock")
                is_session_cmd = any(arg.lower().startswith(sc) or sc in arg.lower() for sc in session_commands)
                try:
                    if is_session_cmd:
                        # Fire and forget — these commands end the session so we can't wait for output
                        _sc.Popen(arg, shell=True)
                        log_entry["result"] = f"command sent: {arg}"
                        log_entry["success"] = True
                    else:
                        result = _sc.run(
                            arg,
                            shell=True,
                            capture_output=True,
                            text=True,
                            timeout=30
                        )
                        log_entry["result"] = result.stdout.strip() or result.stderr.strip() or "done"
                        log_entry["success"] = result.returncode == 0
                except _sc.TimeoutExpired:
                    log_entry["result"] = "command timed out after 30s (may still be running)"
                    log_entry["success"] = True  # treat timeout as sent, not failed
                except Exception as se:
                    log_entry["result"] = f"error: {se}"
                    log_entry["success"] = False

            elif action == "RUN_PYTHON":
                try:
                    exec_globals = {}
                    exec(arg, exec_globals)
                    log_entry["result"] = "success"
                except Exception as pe:
                    log_entry["result"] = f"error: {pe}"
                    log_entry["success"] = False

            elif action == "SCREENSHOT":
                import desktop_actions
                result = desktop_actions.screenshot()
                log_entry["result"] = "screenshot taken" if result == "success" else result
                log_entry["success"] = result == "success"

            elif action.startswith("BROWSER_"):
                # Browser steps are batched and run together at end
                log_entry["result"] = "queued for Puppeteer batch"

        except Exception as e:
            log_entry["result"] = f"exception: {str(e)}"
            log_entry["success"] = False

        steps_log.append(log_entry)

    # ── Step 4: Run batched Puppeteer script if any browser steps ─────────────
    if has_browser_steps and browser_script_path:
        try:
            node_cmd = "node.exe" if os.name == "nt" else "node"
            with open(browser_script_path, "w", encoding="utf-8") as f:
                f.write("\n".join(pup_lines))

            node_result = _sp.run(
                [node_cmd, browser_script_path],
                capture_output=True, text=True, timeout=60,
                cwd=script_dir
            )
            browser_log = {
                "action": "PUPPETEER_BATCH",
                "arg": browser_script_path,
                "result": node_result.stdout.strip() or node_result.stderr.strip() or "done",
                "success": node_result.returncode == 0
            }
            steps_log.append(browser_log)
            try:
                os.remove(browser_script_path)
            except Exception:
                pass
        except FileNotFoundError:
            steps_log.append({
                "action": "PUPPETEER_BATCH", "arg": "",
                "result": "Node.js not found — install Node.js to enable browser automation",
                "success": False
            })
        except Exception as e:
            steps_log.append({
                "action": "PUPPETEER_BATCH", "arg": "",
                "result": f"Puppeteer error: {str(e)}",
                "success": False
            })

    failed = [s for s in steps_log if not s["success"]]
    final_msg = (
        f"Completed {len(steps_log) - len(failed)}/{len(steps_log)} steps successfully."
        if not failed else
        f"{len(steps_log) - len(failed)}/{len(steps_log)} steps succeeded. "
        f"{len(failed)} failed: " + "; ".join(s['action'] for s in failed)
    )

    return len(failed) == 0, steps_log, final_msg, ""



def index_project_directory(directory_path):
    """
    Recursively walk a directory and index all readable text/code files into the doc store.
    Returns (files_indexed: int, total_chunks: int, errors: list).
    """
    SUPPORTED_CODE_EXTENSIONS = {
        '.py', '.js', '.ts', '.tsx', '.jsx', '.html', '.css', '.md',
        '.txt', '.json', '.yaml', '.yml', '.toml', '.ini', '.cfg',
        '.sh', '.bash', '.zsh', '.sql', '.rs', '.go', '.java', '.c',
        '.cpp', '.h', '.hpp', '.rb', '.php', '.swift', '.kt', '.cs',
        '.r', '.scala', '.lua', '.dart', '.vue', '.svelte'
    }
    MAX_FILE_SIZE_BYTES = 1 * 1024 * 1024  # 1MB per file

    files_indexed = 0
    total_chunks = 0
    errors = []
    directory_path = os.path.abspath(directory_path)

    if not os.path.isdir(directory_path):
        return 0, 0, [f"Directory not found: {directory_path}"]

    SKIP_DIRS = {
        '__pycache__', 'node_modules', '.git', '.svn', 'venv',
        'env', '.venv', 'dist', 'build', '.next', '.nuxt', 'coverage'
    }

    for root, dirs, files in os.walk(directory_path):
        # Prune skip directories
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS and not d.startswith('.')]

        for filename in files:
            ext = os.path.splitext(filename)[1].lower()
            if ext not in SUPPORTED_CODE_EXTENSIONS:
                continue

            filepath = os.path.join(root, filename)
            try:
                if os.path.getsize(filepath) > MAX_FILE_SIZE_BYTES:
                    errors.append(f"Skipped (too large): {filepath}")
                    continue

                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()

                if not content.strip():
                    continue

                # Use a relative path as the document name for clarity
                relative_name = os.path.relpath(filepath, directory_path)
                chunks = doc_store.add_document(
                    text=content,
                    filename=relative_name,
                    metadata={
                        'source': 'project_index',
                        'full_path': filepath,
                        'upload_time': datetime.now().isoformat()
                    }
                )
                files_indexed += 1
                total_chunks += chunks

            except Exception as e:
                errors.append(f"Error reading {filepath}: {str(e)}")

    return files_indexed, total_chunks, errors


def handle_tool_call(tool_call_text):
    """
    Sovereign Hybrid Parser: Handles <tool_call> tags AND JSON-leak (addCriterion) formats.
    """
    import re
    import json

    # 1. Try to handle the <tool_call> tags first
    match = re.search(r'<tool_call>\s*(.*?)\s*</tool_call>', tool_call_text, re.IGNORECASE | re.DOTALL)
    if match:
        tool_call_text = match.group(1)

    # 2. Try to handle the 'addCriterion' JSON leak format
    if "addCriterion" in tool_call_text:
        try:
            # Extract the JSON part from the string
            json_match = re.search(r'\{.*\}', tool_call_text, re.DOTALL)
            if json_match:
                data = json.loads(json_match.group(0))
                # Try to extract command. It might be in 'tool_call' or 'action'/'args'
                if 'tool_call' in data:
                    cmd = data['tool_call']
                    if cmd.startswith("run_shell|"):
                        tool_call_text = cmd
                    else:
                        tool_call_text = f"run_shell|{cmd}"
                elif 'action' in data:
                    args = data.get('args', '')
                    tool_call_text = f"run_shell|python desktop_actions.py {data['action']} \"{args}\""
        except Exception:
            pass # Fall back to standard parsing

    # 3. Handle Python-style function calls (e.g., open_app("Spotify"), click(100, 200))
    _KNOWN_FUNCS = ['open_app', 'click', 'type_text', 'press_key', 'screenshot',
                    'click_element_by_name', 'type_into_element', 'open_url']
    func_match = re.match(r'(\w+)\s*\((.*)\)', tool_call_text.strip(), re.DOTALL)
    if func_match and func_match.group(1) in _KNOWN_FUNCS:
        func_name = func_match.group(1)
        raw_args = func_match.group(2).strip()
        # Split arguments by comma and strip quotes
        if raw_args:
            args = [a.strip().strip('"').strip("'") for a in raw_args.split(',')]
            args_str = ' '.join(f'"{a}"' for a in args)
        else:
            args_str = ''
        tool_call_text = f"run_shell|python desktop_actions.py {func_name} {args_str}".strip()

    tool_call_text = tool_call_text.strip()
    first_pipe = tool_call_text.find('|')
    
    if first_pipe == -1:
        tool_name = tool_call_text.lower()
        parts = [tool_name]
    else:
        tool_name = tool_call_text[:first_pipe].strip().lower()
        rest = tool_call_text[first_pipe+1:]
        if tool_name in ['write_file', 'create_docx']:
            parts = [tool_name] + rest.split('|', 1)
        else:
            parts = [tool_name, rest]

    def safe_path(rel_path):
        """Resolve a path and ensure it stays inside the workspace sandbox."""
        abs_path = os.path.realpath(os.path.join(WORKSPACE_DIR, rel_path.strip()))
        workspace_abs = os.path.realpath(WORKSPACE_DIR)
        if not (abs_path.startswith(workspace_abs + os.sep) or abs_path == workspace_abs):
            raise PermissionError(f"Path traversal blocked: {rel_path}")
        return abs_path

    try:
        if tool_name == 'read_file':
            if len(parts) < 2:
                return "[ERROR] read_file requires a file path argument."
            path = safe_path(parts[1])
            if not os.path.exists(path):
                return f"[ERROR] File not found: {parts[1]}"
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return f"[FILE CONTENT: {parts[1]}]\n{content[:5000]}"

        elif tool_name == 'write_file':
            if len(parts) < 3:
                return "[ERROR] write_file requires path and content arguments."
            path = safe_path(parts[1])
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, 'w', encoding='utf-8') as f:
                f.write(parts[2])
            return f"[OK] File written: {parts[1]} ({len(parts[2])} chars)"

        elif tool_name == 'list_dir':
            target = parts[1] if len(parts) > 1 else '.'
            path = safe_path(target)
            if not os.path.isdir(path):
                return f"[ERROR] Not a directory: {target}"
            entries = os.listdir(path)
            return f"[DIR: {target}]\n" + "\n".join(entries)

        elif tool_name == 'run_shell':
            if len(parts) < 2:
                return "[ERROR] run_shell requires a command argument."
            import subprocess
            import sys
            cmd = parts[1]
            if cmd.startswith("python "):
                cmd = f'"{sys.executable}" ' + cmd[7:]
            result = subprocess.run(
                cmd,
                shell=True,
                capture_output=True,
                text=True,
                timeout=30,
                cwd=WORKSPACE_DIR
            )
            out = result.stdout.strip()
            err = result.stderr.strip()
            combined = ""
            if out:
                combined += f"[STDOUT]\n{out}"
            if err:
                combined += f"\n[STDERR]\n{err}"
            combined += f"\n[EXIT CODE] {result.returncode}"
            return combined or "[No output]"

        elif tool_name == 'run_python':
            if len(parts) < 2:
                return "[ERROR] run_python requires code as argument."
            import subprocess, sys
            script_path = os.path.join(WORKSPACE_DIR, f"harness_{uuid.uuid4().hex[:8]}.py")
            try:
                with open(script_path, 'w', encoding='utf-8') as f:
                    f.write(parts[1])
                result = subprocess.run(
                    [sys.executable, script_path],
                    capture_output=True, text=True,
                    timeout=30, cwd=WORKSPACE_DIR
                )
                out = result.stdout.strip()
                err = result.stderr.strip()
                success = result.returncode == 0
                if success:
                    return out or "[OK] Python script executed successfully (no output)."
                else:
                    return f"[ERROR] Python script failed:\n{err}"
            except Exception as e:
                return f"[ERROR] Could not run python script: {str(e)}\n{traceback.format_exc(limit=3)}"
            finally:
                if os.path.exists(script_path):
                    try:
                        os.remove(script_path)
                    except Exception:
                        pass

        elif tool_name == 'open_app':
            if len(parts) < 2:
                return "[ERROR] open_app requires an application name or path."
            import subprocess
            app_target = parts[1].strip()
            try:
                if os.name == 'nt':
                    subprocess.Popen(['start', '', app_target], shell=True)
                else:
                    subprocess.Popen(['xdg-open', app_target])
                return f"[OK] Launched: {app_target}"
            except Exception as e:
                return f"[ERROR] Could not open app: {str(e)}\n{traceback.format_exc(limit=3)}"

        elif tool_name == 'open_url':
            if len(parts) < 2:
                return "[ERROR] open_url requires a URL argument."
            import webbrowser
            url = parts[1].strip()
            webbrowser.open(url)
            return f"[OK] Opened URL: {url}"

        elif tool_name == 'screenshot':
            try:
                import pyautogui
                import io, base64
                img = pyautogui.screenshot()
                buf = io.BytesIO()
                img.save(buf, format='PNG')
                b64 = base64.b64encode(buf.getvalue()).decode()
                return f"[SCREENSHOT] data:image/png;base64,{b64[:200]}... (truncated, {len(b64)} chars total)"
            except Exception as e:
                return f"[ERROR] Screenshot failed: {str(e)}\n{traceback.format_exc(limit=3)}"

        elif tool_name == 'create_docx':
            if len(parts) < 3:
                return "[ERROR] create_docx requires filename and content arguments."
            try:
                from docx import Document as DocxDoc
                path = safe_path(parts[1])
                doc = DocxDoc()
                doc.add_paragraph(parts[2])
                doc.save(path)
                return f"[OK] DOCX created: {parts[1]}"
            except Exception as e:
                return f"[ERROR] create_docx failed: {str(e)}\n{traceback.format_exc(limit=3)}"

        else:
            return f"[ERROR] Unknown tool: {tool_name}"

    except Exception as e:
        return f"[ERROR] Tool execution failed: {str(e)}\n{traceback.format_exc(limit=3)}"


# ═══════════════════════════════════════════════════════════════
#  Routes — Static
# ═══════════════════════════════════════════════════════════════

@app.route('/')
def index():
    """Serve the main interface"""
    try:
        return send_from_directory('.', 'index.html')
    except FileNotFoundError:
        return "<h1>[ERROR] index.html not found</h1><p>Please ensure index.html is in the same directory.</p>", 404


# ═══════════════════════════════════════════════════════════════
#  Routes — Health & Model Info
# ═══════════════════════════════════════════════════════════════

@app.route('/api/health')
def health_check():
    """Health check endpoint — always returns 200"""
    try:
        available_models = check_available_models()
        model_exists = MODEL_NAME in available_models

        return jsonify({
            "status": "connected" if model_exists else "model_not_found",
            "model": MODEL_NAME,
            "available_models": available_models,
            "message": "Ready!" if model_exists else f"Model '{MODEL_NAME}' not found. Pull it with: ollama pull {MODEL_NAME}",
            "rag_stats": doc_store.get_stats(),
            "openclaw_available": True,  # Native agent — always available, no gateway needed
        }), 200

    except requests.exceptions.ConnectionError:
        error_msg = "Ollama is not running. Please start it with: ollama serve"
        print(f"[ERROR] {error_msg}")
        return jsonify({
            "status": "error",
            "message": error_msg,
            "ollama_url": OLLAMA_BASE_URL,
            "rag_stats": doc_store.get_stats(),
            "note": "Server will still work with RAG documents. Start Ollama to enable AI responses.",
            "openclaw_available": True,
        }), 200
    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Health check error: {str(e)}",
            "rag_stats": doc_store.get_stats(),
            "openclaw_available": True,
        }), 200


@app.route('/api/openclaw/status')
def openclaw_status():
    """Native agent status — no external gateway needed."""
    import shutil
    node_available = shutil.which("node.exe" if os.name == "nt" else "node") is not None
    return jsonify({
        "available": True,
        "mode": "native",
        "node_available": node_available,
        "puppeteer_enabled": node_available,
        "message": "Marceline Native Agent is always ready. No gateway required." + (
            " Puppeteer browser automation enabled." if node_available
            else " Install Node.js to enable Puppeteer browser automation."
        ),
        "model": MODEL_NAME,
    })


@app.route('/api/model')
def model_info():
    """Return current model information."""
    return jsonify({
        "model": MODEL_NAME,
        "context_window": MODEL_OPTIONS.get("num_ctx", 4096),
        "vision_supported": is_vision_model(MODEL_NAME)
    })


# ═══════════════════════════════════════════════════════════════
#  Routes — Conversation Management
# ═══════════════════════════════════════════════════════════════

@app.route('/api/conversations', methods=['GET'])
def list_conversations():
    """List all conversations, newest first."""
    try:
        conn = get_db()
        try:
            rows = conn.execute("""
                SELECT c.id, c.title, c.created_at,
                       COUNT(m.id) as message_count
                FROM conversations c
                LEFT JOIN messages m ON m.conversation_id = c.id
                GROUP BY c.id
                ORDER BY c.created_at DESC
            """).fetchall()
        finally:
            conn.close()

        conversations = []
        for row in rows:
            conversations.append({
                "id": row["id"],
                "title": row["title"],
                "message_count": row["message_count"],
                "created_at": row["created_at"],
                "active": False
            })

        return jsonify({"conversations": conversations, "active_id": None})
    except Exception as e:
        print(f"[ERROR] list_conversations: {e}")
        return jsonify({"conversations": [], "active_id": None})


@app.route('/api/conversations', methods=['POST'])
def create_conversation():
    """Create a new conversation."""
    try:
        conv_id = str(uuid.uuid4())
        now = datetime.now().isoformat()

        conn = get_db()
        try:
            conn.execute(
                "INSERT INTO conversations (id, title, created_at) VALUES (?, ?, ?)",
                (conv_id, "New Chat", now)
            )
            conn.commit()
        finally:
            conn.close()

        return jsonify({"id": conv_id, "title": "New Chat"})
    except Exception as e:
        print(f"[ERROR] create_conversation: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/conversations/<conv_id>/messages', methods=['GET'])
def get_conversation_messages(conv_id):
    """Get all messages for a conversation."""
    try:
        conn = get_db()
        try:
            rows = conn.execute(
                "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY created_at ASC",
                (conv_id,)
            ).fetchall()
        finally:
            conn.close()

        messages_list = [{"role": row["role"], "content": row["content"]} for row in rows]
        return jsonify({"messages": messages_list})
    except Exception as e:
        print(f"[ERROR] get_conversation_messages: {e}")
        return jsonify({"messages": []})


@app.route('/api/conversations/<conv_id>/activate', methods=['POST'])
def activate_conversation(conv_id):
    """Activate a conversation (stub)."""
    return jsonify({"success": True})


@app.route('/api/conversations/<conv_id>', methods=['DELETE'])
def delete_conversation(conv_id):
    """Delete a conversation and all its messages."""
    try:
        conn = get_db()
        try:
            conn.execute("DELETE FROM messages WHERE conversation_id = ?", (conv_id,))
            conn.execute("DELETE FROM conversations WHERE id = ?", (conv_id,))
            conn.commit()
        finally:
            conn.close()
        return jsonify({"success": True})
    except Exception as e:
        print(f"[ERROR] delete_conversation: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/conversations/<conv_id>/rename', methods=['POST'])
def rename_conversation(conv_id):
    """Rename a conversation."""
    try:
        data = request.get_json(force=True, silent=True) or {}
        new_title = data.get("title", "").strip()
        if not new_title:
            return jsonify({"error": "Title is required"}), 400

        conn = get_db()
        try:
            conn.execute("UPDATE conversations SET title = ? WHERE id = ?", (new_title, conv_id))
            conn.commit()
        finally:
            conn.close()
        return jsonify({"success": True})
    except Exception as e:
        print(f"[ERROR] rename_conversation: {e}")
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════════════════
#  Routes — Document Upload & RAG
# ═══════════════════════════════════════════════════════════════

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Upload and process documents"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files['file']

        if not file.filename:
            return jsonify({"error": "No file selected"}), 400

        if not allowed_file(file.filename):
            if ALLOWED_EXTENSIONS is None:
                return jsonify({"error": "File must have an extension"}), 400
            else:
                return jsonify({"error": f"File type not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"}), 400

        filename = secure_filename(file.filename)
        if '.' not in filename:
            return jsonify({"error": "Could not determine file extension after sanitization"}), 400
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        file_ext = filename.rsplit('.', 1)[1].lower()

        print(f"[UPLOAD] Processing {filename}...")
        start_time = time.time()

        # Extract text based on file type
        if file_ext == 'pdf':
            text = extract_text_from_pdf(filepath)
        elif file_ext == 'docx':
            text = extract_text_from_docx(filepath)
        elif file_ext == 'zip':
            text = extract_text_from_zip(filepath)
        elif file_ext in ['html', 'htm']:
            text = extract_text_from_html(filepath)
        elif file_ext in ['txt', 'md', 'csv']:
            text = extract_text_from_txt(filepath)
        elif file_ext == 'json':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            text = extract_text_from_txt(filepath)

        if not text.strip():
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({"error": "No text could be extracted"}), 400

        # Add to document store
        chunks_created = doc_store.add_document(
            text=text,
            filename=filename,
            metadata={
                'upload_time': datetime.now().isoformat(),
                'file_size': os.path.getsize(filepath),
                'file_type': file_ext
            }
        )

        processing_time = time.time() - start_time
        print(f"[OK] Processed {filename}: {chunks_created} chunks in {processing_time:.2f}s")

        return jsonify({
            "success": True,
            "message": f"Successfully processed {filename}",
            "filename": filename,
            "chunks_created": chunks_created,
            "processing_time": round(processing_time, 2),
            "file_type": file_ext,
            "stats": doc_store.get_stats()
        })

    except Exception as e:
        if 'filepath' in locals() and os.path.exists(filepath):
            os.remove(filepath)
        print(f"[ERROR] Upload error: {e}")
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500


@app.route('/api/documents', methods=['GET'])
def list_documents():
    """List all documents"""
    return jsonify(doc_store.get_stats())


@app.route('/api/documents/clear', methods=['POST'])
def clear_documents():
    """Clear all documents"""
    doc_store.clear()

    for filename in os.listdir(UPLOAD_FOLDER):
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        try:
            if os.path.isfile(filepath):
                os.unlink(filepath)
        except Exception as e:
            print(f"Error deleting {filepath}: {e}")

    return jsonify({"success": True, "message": "All documents cleared"})


@app.route('/api/documents/<filename>', methods=['DELETE'])
def delete_document(filename):
    """Delete a specific document"""
    filename = secure_filename(filename)
    if not filename:
        return jsonify({"error": "Invalid filename"}), 400

    removed = doc_store.remove_document(filename)

    filepath = os.path.join(UPLOAD_FOLDER, filename)
    if os.path.exists(filepath):
        os.remove(filepath)

    return jsonify({
        "success": True,
        "message": f"Removed {removed} chunks",
        "stats": doc_store.get_stats()
    })


# ═══════════════════════════════════════════════════════════════
#  Routes — Chat (Main Endpoint)
# ═══════════════════════════════════════════════════════════════

@app.route('/api/chat', methods=['POST'])
def chat():
    """Handle chat with multi-turn memory, RAG, think mode, search, and vision."""

    try:
        data = request.get_json(force=True, silent=True)
        if data is None:
            data = {}
        user_message = data.get('message', '').strip()

        if not user_message:
            return jsonify({"error": "No message provided"}), 400

        conversation_id = data.get('conversation_id')
        use_rag = data.get('use_rag', True)
        images = data.get('images', [])

        # ── Detect mode prefixes ──
        active_system_prompt = SYSTEM_PROMPT
        clean_message = user_message
        search_context = ""
        task_description = ""  # Will be set inside [Automate:] block if applicable

        # mode drives which model gets used for the Ollama call further below.
        # "search" is the DEFAULT — it's what gets used when no prefix matches.
        mode = "search"
        model_to_use = SEARCH_MODEL
        model_options_to_use = MODEL_OPTIONS
        harness_task = None  # set when Harness mode should bypass the LLM and call OpenClaw directly

        # Automate mode — Mistral writes a macro, executed locally via pyautogui/pywinauto
        if user_message.startswith("[Automate:"):
            mode = "automate"
            model_to_use = AUTOMATE_MODEL
            model_options_to_use = MODEL_OPTIONS
            task_description = user_message[10:]
            if task_description.endswith("]"):
                task_description = task_description[:-1]
            task_description = task_description.strip()

            clean_message = task_description
            # Ensure ONLY the step-list prompt is sent for reliability if LLM generation runs
            active_system_prompt = (
                "You are an automation planner. The user wants to automate a task on their Windows computer "
                "using pyautogui / pywinauto macros. "
                "Output ONLY a numbered action list. Each line must be: STEP N: ACTION_TYPE | argument. "
                "Valid actions are: LAUNCH_APP (launch an app by executable name), TYPE_TEXT (type text), "
                "PRESS_KEY (key name), OPEN_URL (full URL), WAIT_SECONDS (number), RUN_PYTHON (single line of python). "
                "No explanation. No markdown. No blank lines. "
                "Example for 'open notepad and type hello':\n"
                "STEP 1: LAUNCH_APP | notepad\n"
                "STEP 2: WAIT_SECONDS | 1.5\n"
                "STEP 3: TYPE_TEXT | hello"
            )

            # Turn off RAG context injection for this exact flow
            search_context = ""

        # Harness mode — ANY task the user describes is handed off entirely to
        # OpenClaw ("Open Spotify and play X", "open VS Code and write hello world", etc).
        # No local LLM tool-call loop, no coordinates — OpenClaw drives the OS directly.
        elif user_message.startswith("[Harness:"):
            mode = "harness"
            model_to_use = HARNESS_NARRATION_MODEL
            model_options_to_use = MODEL_OPTIONS
            task = user_message[9:]
            if task.endswith("]"):
                task = task[:-1]
            task = task.strip()
            clean_message = task
            harness_task = task  # signals generate() to call OpenClaw directly instead of Ollama
            active_system_prompt = OPENCLAW_OPERATOR_PROMPT

        # Project mode — index and query a local codebase
        elif user_message.startswith("[Project:"):
            remainder = user_message[9:]
            if remainder.endswith("]"):
                remainder = remainder[:-1]
            remainder = remainder.strip()

            # Check if this is an indexing command (path provided) or a query
            if os.path.isdir(remainder):
                # Index the directory
                files_indexed, total_chunks, errors = index_project_directory(remainder)
                summary = (
                    f"Indexed project at `{remainder}`: "
                    f"{files_indexed} files, {total_chunks} chunks added to RAG store."
                )
                if errors:
                    summary += f"\n\nWarnings ({len(errors)}):\n" + "\n".join(errors[:10])
                return jsonify({"success": True, "message": summary, "stats": doc_store.get_stats()})
            else:
                # Treat it as a project-context query
                clean_message = remainder
                active_system_prompt += (
                    "\n\nThe user is querying their indexed project codebase. "
                    "Answer using the code context provided. Be specific about file names and line logic. "
                    "If you reference code, quote the relevant portion and explain it."
                )

        # Search mode — also reachable as the DEFAULT (no prefix). Mistral, ChatGPT-style.
        elif user_message.startswith("[Search:"):
            mode = "search"
            model_to_use = SEARCH_MODEL
            model_options_to_use = MODEL_OPTIONS
            query = user_message[8:]
            if query.endswith("]"):
                query = query[:-1]
            query = query.strip()
            clean_message = query

            search_results = search_duckduckgo(query)
            if search_results:
                search_context = f"\n=== WEB SEARCH RESULTS FOR: {query} ===\n"
                for i, sr in enumerate(search_results, 1):
                    search_context += f"[{i}] {sr['url']}: {sr['snippet']}\n"
                search_context += "=== END SEARCH RESULTS ===\n"

                active_system_prompt += (
                    "\n\nYou have been given real web search results above. Use them to "
                    "answer accurately. Cite which source number you are drawing from."
                )

        # Think mode — DeepSeek reasoning model, shows its work
        elif user_message.startswith("[Think:"):
            mode = "think"
            model_to_use = THINK_MODEL
            model_options_to_use = THINK_MODEL_OPTIONS
            clean_message = user_message[7:]
            if clean_message.endswith("]"):
                clean_message = clean_message[:-1]
            clean_message = clean_message.strip()
            active_system_prompt += (
                "\n\nThe user has activated Think Mode. Before giving your final answer, "
                "reason through the problem step by step inside <thinking> tags. Show your "
                "work, consider edge cases, and think out loud. Then give your clean final "
                "answer after the closing </thinking> tag."
            )

        elif user_message.startswith("[Plan:"):
            mode = "plan"
            model_to_use = SEARCH_MODEL
            model_options_to_use = MODEL_OPTIONS
            task = user_message[6:]
            if task.endswith("]"):
                task = task[:-1]
            task = task.strip()
            clean_message = task
            active_system_prompt += (
                "\n\nThe user has activated Planner Mode. Your job is to break this task into a "
                "clear numbered step-by-step execution plan. Format your response as:\n"
                "PLAN:\n1. Step one\n2. Step two\n...\n\n"
                "TOOLS NEEDED:\n- List each tool_call you will use, in order\n\n"
                "ESTIMATED STEPS: N\n\n"
                "Do not execute anything yet. Only produce the plan."
            )

        # ── Ensure conversation exists ──
        created_new_conversation = False
        if not conversation_id:
            conversation_id = str(uuid.uuid4())
            now = datetime.now().isoformat()
            conn = get_db()
            try:
                conn.execute(
                    "INSERT INTO conversations (id, title, created_at) VALUES (?, ?, ?)",
                    (conversation_id, "New Chat", now)
                )
                conn.commit()
            finally:
                conn.close()
            created_new_conversation = True

        # ── Load conversation history from DB ──
        db_history = []
        try:
            conn = get_db()
            try:
                # Limit to the last 40 messages to prevent exceeding the context window
                rows = conn.execute(
                    """
                    SELECT role, content FROM (
                        SELECT role, content, created_at 
                        FROM messages 
                        WHERE conversation_id = ? 
                        ORDER BY created_at DESC 
                        LIMIT 40
                    ) ORDER BY created_at ASC
                    """,
                    (conversation_id,)
                ).fetchall()
            finally:
                conn.close()
            db_history = [{"role": row["role"], "content": row["content"]} for row in rows]
        except Exception as e:
            print(f"[WARN] Could not load history: {e}")

        # ── RAG context ──
        rag_context = ""
        sources = []
        if use_rag and doc_store.documents:
            rag_start = time.time()
            results = doc_store.search(clean_message, top_k=3, min_score=0.15)
            rag_time = time.time() - rag_start
            print(f"[SEARCH] RAG search: {len(results)} results in {rag_time:.2f}s")

            if results:
                rag_context = "\n=== CONTEXT FROM DOCUMENTS ===\n"
                for i, result in enumerate(results, 1):
                    doc = result['document']
                    rag_context += f"\n[Source {i} - {doc['filename']}]:\n{doc['text'][:600]}\n"
                    sources.append({
                        'filename': doc['filename'],
                        'chunk_index': doc['chunk_index'],
                        'relevance_score': round(result['score'], 3),
                        'chunk_type': doc.get('chunk_type', 'paragraph')
                    })
                rag_context += "\n=== END CONTEXT ===\n"

        # ── Build system message with context ──
        full_system = active_system_prompt
        if rag_context:
            full_system += f"\n\n{rag_context}"
        if search_context:
            full_system += f"\n\n{search_context}"

        # ── Build messages array for Ollama /api/chat ──
        ollama_messages = [{"role": "system", "content": full_system}]
        ollama_messages.extend(db_history)

        # Add the new user message (with images if any)
        user_msg_obj = {"role": "user", "content": clean_message}
        if images:
            user_msg_obj["images"] = images
        ollama_messages.append(user_msg_obj)

        print(f"[QUERY] Query: {clean_message[:100]}...")
        print(f"[QUERY] Conversation: {conversation_id[:8]}... | History: {len(db_history)} msgs")

        def generate():
            try:
                # ── Harness mode: hand the ENTIRE task straight to OpenClaw. ──
                # No local LLM tool-call loop. "Open Spotify and play X",
                # "open VS Code and write hello world" — OpenClaw drives the OS directly.
                if harness_task:
                    yield f"data: {json.dumps({'type': 'chunk', 'content': '**🦾 Harness Mode (OpenClaw)** — executing: ' + harness_task + chr(10) + chr(10)})}\n\n"
                    try:
                        import openclaw_bridge
                        if not openclaw_bridge.check_deps():
                            yield f"data: {json.dumps({'type': 'chunk', 'content': '⚠️ OpenClaw (npx openclaw) was not found on PATH. Install it with `npm install -g openclaw` and try again.' + chr(10)})}\n\n"
                        else:
                            result = openclaw_bridge.run_primitive(harness_task)
                            if result == "success":
                                yield f"data: {json.dumps({'type': 'chunk', 'content': '✅ Task executed successfully.' + chr(10)})}\n\n"
                            else:
                                yield f"data: {json.dumps({'type': 'chunk', 'content': '⚠️ Could not complete task: ' + str(result)[:500] + chr(10)})}\n\n"
                            full_response = f"[Harness/OpenClaw] {harness_task} -> {result}"
                            add_to_history(clean_message, full_response)
                    except Exception as e:
                        yield f"data: {json.dumps({'type': 'chunk', 'content': '❌ OpenClaw execution error: ' + str(e) + chr(10)})}\n\n"

                    if created_new_conversation:
                        t = threading.Thread(target=generate_title_async, args=(conversation_id, clean_message), daemon=True)
                        t.start()
                    yield f"data: {json.dumps({'type': 'done', 'conversation_id': conversation_id})}\n\n"
                    return

                # Send sources first if we have any
                if sources:
                    yield f"data: {json.dumps({'type': 'sources', 'sources': sources})}\n\n"

                # Vision overrides the mode's normal model when images are attached
                effective_model = VISION_MODEL if images else model_to_use

                # Check if user sent images to a non-vision model
                if images and not is_vision_model(effective_model):
                    yield f"data: {json.dumps({'type': 'error', 'content': 'The current model does not support image understanding. Switch to a vision model like qwen2.5vl:latest in server.py to analyze images.'})}\n\n"
                    return

                ollama_request = {
                    "model": effective_model,
                    "messages": ollama_messages,
                    "stream": True,
                    "options": model_options_to_use,
                    "keep_alive": GPU_KEEP_ALIVE
                }

                if images:
                    # Yield an initial "thinking" chunk to keep the frontend connection alive
                    yield f"data: {json.dumps({'type': 'chunk', 'content': '*(Processing vision...)* '})}\n\n"

                response = requests.post(
                    f"{OLLAMA_BASE_URL}/api/chat",
                    json=ollama_request,
                    headers={"Content-Type": "application/json"},
                    stream=True,
                    timeout=180
                )

                if response.status_code != 200:
                    yield f"data: {json.dumps({'type': 'error', 'content': f'Ollama API error: HTTP {response.status_code}'})}\n\n"
                    return

                full_response = ""
                for line in response.iter_lines():
                    if line:
                        try:
                            chunk = json.loads(line)
                        except json.JSONDecodeError as e:
                            print(f"[WARN] Skipping non-JSON line from Ollama: {line!r} — {e}")
                            continue

                        # Ollama /api/chat format: chunk["message"]["content"]
                        text_chunk = chunk.get("message", {}).get("content", "")
                        full_response += text_chunk
                        yield f"data: {json.dumps({'type': 'chunk', 'content': text_chunk})}\n\n"

                        if chunk.get("done", False):
                            break

                # Save to in-memory history
                add_to_history(clean_message, full_response)

                # Save to SQLite
                try:
                    now = datetime.now().isoformat()
                    with db_lock:
                        conn = get_db()
                        try:
                            # Save user message
                            conn.execute(
                                "INSERT INTO messages (id, conversation_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
                                (str(uuid.uuid4()), conversation_id, "user", clean_message, now)
                            )
                            # Save assistant response
                            conn.execute(
                                "INSERT INTO messages (id, conversation_id, role, content, created_at) VALUES (?, ?, ?, ?, ?)",
                                (str(uuid.uuid4()), conversation_id, "assistant", full_response, now)
                            )
                            conn.commit()
                        finally:
                            conn.close()
                except Exception as db_err:
                    print(f"[ERROR] DB save failed: {db_err}")

                # -- AUTO-EXTRACT AND EXECUTE TOOL CALLS --
                # This catches BOTH tagged <tool_call>run_shell|...</tool_call>
                # AND untagged plain-text run_shell|... commands from the LLM.
                import re as _re
                tool_commands = []

                # Priority 1: Extract properly tagged tool calls
                tagged = _re.findall(r'<tool_call>\s*(.*?)\s*</tool_call>', full_response, _re.IGNORECASE | _re.DOTALL)
                for cmd in tagged:
                    cmd = cmd.strip()
                    if cmd and cmd not in tool_commands:
                        tool_commands.append(cmd)

                # Priority 2: If NO tagged commands found, try to extract untagged ones
                # Matches patterns like: run_shell|python desktop_actions.py ...
                if not tool_commands:
                    untagged = _re.findall(r'(run_shell\|[^\n<`*]+)', full_response)
                    for cmd in untagged:
                        cmd = cmd.strip()
                        if cmd and cmd not in tool_commands:
                            tool_commands.append(cmd)

                # Priority 3: Extract Qwen JSON leak (addCriterion)
                if not tool_commands:
                    json_leaks = _re.findall(r'(addCriterion\s*\{.*?\})', full_response, _re.DOTALL)
                    for leak in json_leaks:
                        leak = leak.strip()
                        if leak and leak not in tool_commands:
                            tool_commands.append(leak)

                # Priority 4: Extract raw Python function calls like open_app("Spotify"), click(100, 200), etc.
                # The AI sometimes outputs bare function calls instead of using tags or run_shell
                if not tool_commands:
                    _KNOWN_VERBS = ['open_app', 'click', 'type_text', 'press_key', 'screenshot',
                                    'click_element_by_name', 'type_into_element', 'open_url']
                    for verb in _KNOWN_VERBS:
                        # Match verb("arg") or verb('arg') or verb(arg1, arg2)
                        pattern = verb + r'\s*\(\s*["\']?([^)]*?)["\']?\s*\)'
                        matches = _re.findall(pattern, full_response)
                        for m in matches:
                            m = m.strip().strip("'\"")
                            if verb == 'screenshot':
                                converted = f'run_shell|python desktop_actions.py screenshot'
                            elif verb == 'click' and ',' in m:
                                # click(100, 200) -> click 100 200
                                coords = [c.strip().strip("'\"") for c in m.split(',')]
                                converted = f'run_shell|python desktop_actions.py click {coords[0]} {coords[1]}'
                            elif verb in ('click_element_by_name', 'type_into_element') and ',' in m:
                                parts = [p.strip().strip("'\"") for p in m.split(',', 1)]
                                converted = f'run_shell|python desktop_actions.py {verb} "{parts[0]}" "{parts[1]}"'
                            else:
                                converted = f'run_shell|python desktop_actions.py {verb} "{m}"'
                            if converted not in tool_commands:
                                tool_commands.append(converted)

                if tool_commands:
                    yield f"data: {json.dumps({'type': 'chunk', 'content': chr(10)*2 + '---' + chr(10) + '**🔧 Executing ' + str(len(tool_commands)) + ' tool command(s)...**' + chr(10)})}\n\n"
                    for i, tc in enumerate(tool_commands, 1):
                        try:
                            result = handle_tool_call(tc)
                            icon = '✅' if 'error' not in str(result).lower()[:50] else '⚠️'
                            yield f"data: {json.dumps({'type': 'chunk', 'content': f'{icon} **Step {i}**: `{tc[:80]}` → {str(result)[:200]}' + chr(10)})}\n\n"
                        except Exception as e:
                            yield f"data: {json.dumps({'type': 'chunk', 'content': f'❌ **Step {i} FAILED**: `{tc[:80]}` → {str(e)[:200]}' + chr(10)})}\n\n"

                # -- LEGACY AUTOMATION EXECUTION (for [Automate:] prefix) --
                if user_message.startswith("[Automate:"):
                    yield f"data: {json.dumps({'type': 'chunk', 'content': chr(10)*2 + '---' + chr(10) + '*Executing automation script...*'})}\n\n"
                    
                    try:
                        step_generator = execute_step_list(full_response)
                        success_any = False
                        for step_output in step_generator:
                            yield f"data: {json.dumps({'type': 'chunk', 'content': step_output})}\n\n"
                            if "✅ Completed" in step_output:
                                success_any = True
                                
                        if not success_any:
                            msg = chr(10) + "**⚠️ LLM script failed. Trying vision-guided control...**" + chr(10)
                            yield f"data: {json.dumps({'type': 'chunk', 'content': msg})}\n\n"
                            for vision_status in run_vision_guided_automation(task_description):
                                yield f"data: {json.dumps({'type': 'chunk', 'content': vision_status})}\n\n"
                                
                        log_automation(task_description or clean_message, full_response[:300], "Step execution complete.")
                    except Exception as e:
                        err_msg = f"Automation Execution Error: {str(e)}"
                        yield f"data: {json.dumps({'type': 'chunk', 'content': chr(10)*2 + '**[ERROR]**' + chr(10) + '```' + chr(10) + err_msg + chr(10) + '```'})}\n\n"
                        log_automation(task_description or clean_message, full_response[:300], err_msg)

                # Generate title in background if this is a new conversation
                if created_new_conversation:
                    t = threading.Thread(
                        target=generate_title_async,
                        args=(conversation_id, clean_message, effective_model),
                        daemon=True
                    )
                    t.start()

                yield f"data: {json.dumps({'type': 'done', 'conversation_id': conversation_id})}\n\n"

            except requests.exceptions.ConnectionError:
                error_msg = "Ollama is not running. Please start it with: ollama serve"
                print(f"[ERROR] {error_msg}")
                yield f"data: {json.dumps({'type': 'error', 'content': error_msg})}\n\n"
            except Exception as e:
                print(f"[ERROR] Stream error: {e}")
                yield f"data: {json.dumps({'type': 'error', 'content': f'Server error: {str(e)}'})}\n\n"

        return Response(stream_with_context(generate()), mimetype='text/event-stream')

    except Exception as e:
        print(f"[ERROR] Chat error: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500


# ═══════════════════════════════════════════════════════════════
#  Routes — Misc
# ═══════════════════════════════════════════════════════════════

@app.route('/debug')
def debug():
    """Debug endpoint"""
    return jsonify({
        "ollama_url": OLLAMA_BASE_URL,
        "model": MODEL_NAME,
        "upload_folder": UPLOAD_FOLDER,
        "vector_store_path": VECTOR_STORE_PATH,
        "rag_stats": doc_store.get_stats(),
        "conversation_count": get_conversation_count(),
        "timestamp": datetime.now().isoformat()
    })


@app.route('/api/chat/clear', methods=['POST'])
def clear_chat_history():
    """Clear in-memory conversation history"""
    with history_lock:
        conversation_history.clear()
    return jsonify({"success": True, "message": "Chat history cleared"})


@app.route('/api/chat/history', methods=['GET'])
def get_chat_history():
    """Get in-memory conversation history"""
    with history_lock:
        history_copy = list(conversation_history)
    return jsonify({
        "history": history_copy,
        "message_count": len(history_copy) // 2
    })


@app.route('/api/tool', methods=['POST'])
def run_tool():
    """Endpoint for the frontend to execute a tool call from the LLM."""
    try:
        data = request.get_json()
        tool_call = data.get('tool_call', '')
        if not tool_call:
            return jsonify({"error": "No tool call provided"}), 400

        result = handle_tool_call(tool_call)
        return jsonify({"result": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/workspace', methods=['GET'])
def list_workspace():
    """List all files in the workspace directory recursively."""
    try:
        if not os.path.exists(WORKSPACE_DIR):
            os.makedirs(WORKSPACE_DIR, exist_ok=True)
            
        tree = []
        for root, dirs, files in os.walk(WORKSPACE_DIR):
            # Prune hidden dirs
            dirs[:] = [d for d in dirs if not d.startswith('.')]
            
            rel_path = os.path.relpath(root, WORKSPACE_DIR)
            if rel_path == '.':
                rel_path = ''
                
            for file in files:
                if file.startswith('.'):
                    continue
                file_rel = os.path.join(rel_path, file).replace('\\', '/')
                file_abs = os.path.join(root, file)
                tree.append({
                    "path": file_rel,
                    "name": file,
                    "size": os.path.getsize(file_abs),
                    "modified": os.path.getmtime(file_abs)
                })
                
        return jsonify({"files": tree})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/workspace/<path:filename>', methods=['GET'])
def get_workspace_file(filename):
    """Serve a file from the workspace."""
    try:
        # Prevent path traversal outside WORKSPACE_DIR
        safe_path = os.path.abspath(os.path.join(WORKSPACE_DIR, filename))
        workspace_abs = os.path.abspath(WORKSPACE_DIR)
        if not (safe_path.startswith(workspace_abs + os.sep) or safe_path == workspace_abs):
            return jsonify({"error": "Access denied"}), 403
            
        return send_from_directory(WORKSPACE_DIR, filename)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════════════════
#  STARTUP
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    # Initialize database
    init_db()

    # Dynamically find an open port
    import socket
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    s.bind(('127.0.0.1', 0))
    port = s.getsockname()[1]
    s.close()
    
    local_ip = get_local_ip()

    print("======================================================================")
    print("  MARCELINE - Local AI Assistant")
    print("=" * 70)
    print(f"  Server:        http://{local_ip}:{port}")
    print(f"  Search model:  {SEARCH_MODEL}  (default mode, ChatGPT-style)")
    print(f"  Think model:   {THINK_MODEL}  (reasoning)")
    print(f"  Automate model:{AUTOMATE_MODEL}  (pyautogui/pywinauto macro planner)")
    print(f"  Harness:       OpenClaw (no LLM in the loop)")
    print(f"  Vision model:  {VISION_MODEL}")
    print(f"  Database:      {DB_PATH}")

    conv_count = get_conversation_count()
    print(f"  Conversations: {conv_count}")

    stats = doc_store.get_stats()
    print(f"  Documents:     {stats['total_files']} files, {stats['total_chunks']} chunks")

    if stats['files']:
        print("  Files in store:")
        for file_info in stats['files']:
            print(f"    - {file_info['name']} ({file_info['chunks']} chunks)")

    import shutil
    node_path = shutil.which("node.exe") if os.name == "nt" else shutil.which("node")
    print("\n[CHECK] Checking Puppeteer/Node.js...")
    if node_path:
        print(f"  [OK] Node.js found at: {node_path}")
        print(f"  [OK] Browser automation (Puppeteer) is ENABLED")
        print(f"       If puppeteer module is missing, run: npm install puppeteer")
    else:
        print(f"  [INFO] Node.js not found. Browser automation (Puppeteer) is disabled.")
        print(f"         Install Node.js from https://nodejs.org to enable it.")
        print(f"  [OK] Desktop automation (pyautogui) is always available.")

    print("\n[CHECK] Checking Ollama...")
    if is_ollama_reachable():
        models = check_available_models()
        for needed in [SEARCH_MODEL, THINK_MODEL, VISION_MODEL]:
            if needed in models:
                print(f"  [OK] '{needed}' found")
            else:
                print(f"  [WARN] '{needed}' not found — pull it with: ollama pull {needed}")
    else:
        print(f"  [FAIL] Ollama is not reachable at {OLLAMA_BASE_URL}")
        print("    Start it with: ollama serve")



    print("\n" + "=" * 70 + "\n")


    run_desktop = '--desktop' in sys.argv

    def run_flask():
        app.run(host='0.0.0.0', port=port, debug=False, use_reloader=False)

    if run_desktop:
        print("  Starting in Desktop mode...")
        try:
            import webview
            
            # Start Flask in a background thread
            flask_thread = threading.Thread(target=run_flask)
            flask_thread.daemon = True
            flask_thread.start()
            
            # Wait for server to be ready
            time.sleep(1)
            
            # Launch the webview window
            webview.create_window(
                title='Marceline', 
                url=f'http://127.0.0.1:{port}', 
                width=1200, 
                height=800, 
                background_color='#1a0b2e', 
                min_size=(800, 600)
            )
            webview.start(gui='edgechromium' if os.name == 'nt' else None)
            
        except ImportError:
            print("  [ERROR] pywebview not installed. Cannot start desktop mode.")
            print("  Falling back to normal web server mode...")
            run_flask()
    else:
        print("  Starting in Web Server mode (use --desktop for standalone app)")
        run_flask()